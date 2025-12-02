"""
Document Compare Launch Content Generator
AI-powered marketing content generation for Genie AI's Document Compare feature
"""

import streamlit as st
from openai import OpenAI
import openai as openai_legacy
try:
    import requests
except ImportError:
    requests = None
from datetime import datetime
import json
import io
from typing import Optional
import re

# PDF and DOCX parsing
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


# ============================================================================
# CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="Document Compare Launch Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for beautiful UI
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
    
    /* Override default Streamlit fonts */
    .stApp {
        font-family: 'Space Grotesk', sans-serif;
    }
    
    /* Main container styling */
    .main-header {
        background: linear-gradient(135deg, #673ab7 0%, #007aff 100%);
        padding: 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        color: white;
        text-align: center;
    }
    
    .main-header h1 {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        background: linear-gradient(90deg, #ffffff, #dfe7ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .main-header p {
        color: #e5e7f3;
        font-size: 1.1rem;
    }
    
    /* Step indicator */
    .step-container {
        display: flex;
        justify-content: space-between;
        margin-bottom: 2rem;
        padding: 1rem;
        background: #f8f9fa;
        border-radius: 12px;
    }
    
    .step {
        display: flex;
        flex-direction: column;
        align-items: center;
        flex: 1;
        position: relative;
    }
    
    .step-number {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 600;
        margin-bottom: 0.5rem;
    }
    
    .step-active .step-number {
        background: linear-gradient(135deg, #673ab7, #007aff);
        color: white;
    }
    
    .step-completed .step-number {
        background: #10b981;
        color: white;
    }
    
    .step-pending .step-number {
        background: #e5e7eb;
        color: #6b7280;
    }
    
    /* Card styling */
    .content-card {
        background: white;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .content-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        border-color: #673ab7;
    }
    
    /* Quote card */
    .quote-card {
        background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%);
        border-left: 4px solid #f59e0b;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        font-style: italic;
    }
    
    /* Draft version card */
    .draft-card {
        background: #f8fafc;
        border: 2px solid #e2e8f0;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        transition: all 0.2s ease;
    }
    
    .draft-card.selected {
        border-color: #e94560;
        background: #fff5f7;
    }
    
    .draft-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #e2e8f0;
    }
    
    .version-badge {
        background: #673ab7;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
    }
    
    .timestamp {
        color: #64748b;
        font-size: 0.85rem;
    }
    
    /* Approval checkbox styling */
    .approval-section {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        border: 2px solid #10b981;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    /* Button styling */
    .stButton > button {
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 500;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        transition: all 0.2s ease;
    }
    
    /* Progress bar */
    .progress-container {
        background: #e5e7eb;
        border-radius: 10px;
        height: 8px;
        margin: 1rem 0;
        overflow: hidden;
    }
    
    .progress-bar {
        background: linear-gradient(90deg, #673ab7, #007aff);
        height: 100%;
        border-radius: 10px;
        transition: width 0.3s ease;
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background: #f8f9fa;
    }
    
    /* File uploader */
    .uploadedFile {
        background: #f0f9ff;
        border: 1px solid #0ea5e9;
        border-radius: 8px;
    }
    
    /* Chat message styling */
    .chat-message {
        padding: 1rem;
        border-radius: 12px;
        margin: 0.5rem 0;
    }
    
    .chat-assistant {
        background: #f1f5f9;
        border-left: 3px solid #673ab7;
    }
    
    .chat-user {
        background: #eff6ff;
        border-left: 3px solid #007aff;
    }
</style>
""", unsafe_allow_html=True)

LAUNCH_REQUIREMENTS = """- Repeatable and sustainable workflow
- Minimize human intervention time/actions
- Tone must fit founders/executives at 4-50 person US businesses
- Channel-appropriate formatting (LinkedIn, newsletter, blog)
- Drive excitement/engagement for the new feature
- Do not rely heavily on existing channel content; keep ideas fresh
- Human QA is expected; keep automation efficient (< 2 hours build/run)
- Final marketing outputs must be provided in a labeled doc"""

AUDIENCE_PROFILE = """- Executives in 4-50 person businesses
- Not legally trained
- Want to minimize legal time and spend
- Need reliability and accessibility
- US-based"""

# Model routing (use broadly available aliases first)
MODEL_GENERATE = "gpt-4o-mini"
MODEL_QUOTE = "gpt-4o-mini"
MODEL_FALLBACKS = ["gpt-5", "o4-mini-2025-04-16"]
MODEL_TEMP_RESTRICTED = {"gpt-5", "o4-mini-2025-04-16", "o4-mini"}
MODEL_OPTIONS = ["gpt-4o-mini", "gpt-4o", "gpt-4.1", "gpt-4.1-nano", "gpt-5", "o4-mini-2025-04-16"]


# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    """Initialize all session state variables"""
    defaults = {
        # Navigation
        'current_step': 1,
        
        # Step 1: Uploaded files
        'uploaded_files': [],
        'extracted_texts': {},
        
        # Step 2: Content requirements
        'content_type': None,
        'platform': None,
        'content_length': None,
        'content_bundle': ["📝 Blog Post", "📱 LinkedIn Post", "📧 Newsletter Email"],
        'length_preferences': {
            'blog': "900-1200 words (standard launch story)",
            'linkedin': "120-180 words (standard post)",
            'newsletter': "280-400 words (feature highlight)"
        },
        'feature_focus': 'Document Compare',  # Pre-set for this specific task
        'additional_context': '',
        'openai_force_legacy': False,
        'openai_use_rest': True,  # avoid httpx/proxies issues by default
        'model_choice': MODEL_GENERATE,
        
        # Step 3: Tone & Brand
        'tone': None,
        'brand_guidelines': """- Genie AI brand: trust, clarity, simplicity; build subconscious trust for legal/sensitive use-cases
- Speak to US founders/executives (4-50 ppl), non-lawyers
- Focus on time saved, clarity, and reliability; avoid legal jargon
- Confident, concise, trustworthy; avoid hype and clichés
- Lead with outcomes (faster reviews, fewer mistakes), back with specifics
- No emojis unless explicitly requested
- End with a clear CTA to try Document Compare""",
        'example_content': '',
        
        # Step 4: Quotes
        'extracted_quotes': [],
        'selected_quotes': [],
        
        # Step 5: Drafts & Versions
        'drafts': [],  # List of draft versions
        'current_draft_index': 0,
        'current_draft_index_by_type': {},
        'draft_approved': False,
        'feedback_chat': {
            '📝 Blog Post': [],
            '📱 LinkedIn Post': [],
            '📧 Newsletter Email': []
        },
        
        # Step 6: Final content
        'final_contents': {},  # Dict of final content pieces by draft_id
        'generation_complete': False,
        
        # Chat history for requirements gathering
        'chat_history': [],
        'requirements_complete': False,
        
        # OpenAI client
        'openai_client': None,
        'debug_logs': [],
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()


# ============================================================================
# SOURCE MATERIAL LOADING
# ============================================================================

def load_source_materials():
    """Load and cache all Document Compare source materials"""
    if 'source_materials_loaded' not in st.session_state:
        source_texts = {}

        # Load Product Roadmap Summary
        try:
            from docx import Document
            doc = Document('task-feature-information/TASK_ Product Roadmap Summary.docx')
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            source_texts['product_roadmap'] = text
        except Exception as e:
            source_texts['product_roadmap'] = f"Error loading roadmap: {e}"

        # Load Customer Feedback
        try:
            doc = Document('task-feature-information/TASK_ Customer feedback snippets.docx')
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            source_texts['customer_feedback'] = text
        except Exception as e:
            source_texts['customer_feedback'] = f"Error loading feedback: {e}"

        # Load Engineering Ticket
        try:
            doc = Document('task-feature-information/TASK_ Linear Ticket - Engineering.docx')
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            source_texts['engineering_ticket'] = text
        except Exception as e:
            source_texts['engineering_ticket'] = f"Error loading ticket: {e}"

        # Load Meeting Transcript
        try:
            doc = Document('task-feature-information/TASK_ Meeting Transcript (Product + Engineering + Marketing).docx')
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            source_texts['meeting_transcript'] = text
        except Exception as e:
            source_texts['meeting_transcript'] = f"Error loading transcript: {e}"

        # Load Meeting Notes PDF
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader('task-feature-information/TASK_ Marketing & Product Meeting Notes.pdf')
            text = ''
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            source_texts['meeting_notes'] = text
        except Exception as e:
            source_texts['meeting_notes'] = f"Error loading meeting notes: {e}"

        st.session_state.source_materials = source_texts
        st.session_state.extracted_texts.update(source_texts)
        st.session_state.source_materials_loaded = True


# Load source materials on startup
load_source_materials()


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def log_debug(message: str):
    """Append a timestamped debug message to session logs."""
    if 'debug_logs' not in st.session_state:
        st.session_state.debug_logs = []
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.debug_logs.append(f"[{timestamp}] {message}")
    # Keep log from growing indefinitely
    st.session_state.debug_logs = st.session_state.debug_logs[-200:]


def get_model_choice() -> str:
    """Return current selected model."""
    return st.session_state.get('model_choice', MODEL_GENERATE)


def get_feedback_notes(content_type: str) -> str:
    """Aggregate feedback chat for a given content type."""
    if 'feedback_chat' not in st.session_state:
        return ""
    feedback = st.session_state.feedback_chat.get(content_type, [])
    if not feedback:
        return ""
    return "\n".join([f"{m['role']}: {m['content']}" for m in feedback])


def get_openai_client() -> Optional[OpenAI]:
    """Get or create OpenAI client"""
    if st.session_state.get('openai_use_rest'):
        if not st.session_state.get('openai_api_key'):
            log_debug("No API key set; cannot use REST mode")
            return None
        log_debug("Using REST mode for OpenAI")
        st.session_state.openai_client = "rest"
        return st.session_state.openai_client

    if st.session_state.openai_client is None:
        api_key = st.session_state.get('openai_api_key')
        if api_key:
            if st.session_state.get('openai_force_legacy'):
                openai_legacy.api_key = api_key
                st.session_state.openai_client = openai_legacy
            else:
                try:
                    st.session_state.openai_client = OpenAI(api_key=api_key)
                    log_debug("Initialized OpenAI client (httpx)")
                    st.session_state.openai_error = None
                except Exception as e:  # capture TypeError and any transport proxy errors
                    openai_legacy.api_key = api_key
                    st.session_state.openai_client = openai_legacy
                    st.session_state.openai_error = f"OpenAI client fallback in use: {e}"
                    st.session_state.openai_force_legacy = True
                    log_debug(f"Switched to legacy OpenAI client: {e}")
    return st.session_state.openai_client


def chat_completion(client, **kwargs):
    """Compatibility wrapper for new and legacy OpenAI clients"""
    if st.session_state.get("openai_use_rest") or client == "rest":
        log_debug("Chat completion via REST fallback")
        return rest_chat_completion(**kwargs)

    # If we already know there's a proxies issue, go straight to REST fallback
    if st.session_state.get("openai_error") and "proxies" in st.session_state.openai_error.lower():
        log_debug("Chat completion proxied to REST due to previous proxy error")
        return rest_chat_completion(**kwargs)

    try:
        if hasattr(client, "chat") and hasattr(client.chat, "completions"):
            return client.chat.completions.create(**kwargs)
        if hasattr(client, "ChatCompletion"):
            # response_format not supported in legacy SDK
            kwargs.pop("response_format", None)
            return client.ChatCompletion.create(**kwargs)
    except Exception as e:
        st.session_state.openai_error = f"OpenAI chat error: {e}"
        log_debug(f"Chat completion error, switching to REST: {e}")
        return rest_chat_completion(**kwargs)

    raise RuntimeError("OpenAI client is not initialized correctly.")


def rest_chat_completion(**kwargs):
    """HTTP fallback to avoid httpx/proxies issues."""
    api_key = st.session_state.get("openai_api_key")
    if not api_key:
        raise RuntimeError("Missing OpenAI API key.")
    if requests is None:
        raise RuntimeError("requests not installed. Run: pip install requests")

    payload = {
        "model": kwargs.get("model", get_model_choice()),
        "messages": kwargs.get("messages", []),
        "temperature": kwargs.get("temperature", 0.7),
    }
    if kwargs.get("response_format"):
        payload["response_format"] = kwargs["response_format"]

    def _call(model_name):
        log_debug(f"REST chat call: model={model_name}, messages={len(payload['messages'])}")
        body = dict(payload)
        body["model"] = model_name
        if model_name in MODEL_TEMP_RESTRICTED and "temperature" in body:
            body.pop("temperature", None)
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=body,
            timeout=600,
        )
        resp.raise_for_status()
        data = resp.json()
        content = data["choices"][0]["message"]["content"]
        class _Message:
            def __init__(self, c):
                self.content = c
        class _Choice:
            def __init__(self, c):
                self.message = _Message(c)
        class _Response:
            def __init__(self, c):
                self.choices = [_Choice(c)]
        return _Response(content)

    try:
        result = _call(payload["model"])
        st.session_state.openai_error = None
        log_debug("REST chat call succeeded")
        return result
    except requests.HTTPError as e:
        status = e.response.status_code if hasattr(e, "response") and e.response is not None else None
        log_debug(f"REST chat call failed ({status}): {e}")
        if status in (400, 404):
            for fb in MODEL_FALLBACKS:
                log_debug(f"Falling back to {fb}")
                try:
                    result = _call(fb)
                    st.session_state.openai_error = None
                    log_debug("Fallback REST chat call succeeded")
                    return result
                except Exception as e2:
                    detail = ""
                    if hasattr(e2, "response") and e2.response is not None:
                        try:
                            detail = e2.response.text
                        except Exception:
                            detail = ""
                    log_debug(f"Fallback REST chat call failed: {e2} {detail}")
                    continue
            st.session_state.openai_error = f"REST fallback failed: {e}"
            raise
        st.session_state.openai_error = f"REST fallback failed: {e}"
        raise


def extract_text_from_file(uploaded_file) -> str:
    """Extract text from uploaded file (PDF, DOCX, or TXT)"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    file_content = uploaded_file.getvalue()
    
    try:
        if file_type == 'txt':
            # Try multiple encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return file_content.decode('utf-8', errors='ignore')
        
        elif file_type == 'pdf':
            if PdfReader is None:
                return "[Error: PyPDF2 not installed. Run: pip install PyPDF2]"
            pdf_reader = PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text if text.strip() else "[PDF extracted but no text found - may be image-based PDF]"
        
        elif file_type in ['docx', 'doc']:
            if Document is None:
                return "[Error: python-docx not installed. Run: pip install python-docx]"
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text if text.strip() else "[DOCX extracted but no text found]"
        
        else:
            return f"[Unable to extract text from .{file_type} file - unsupported format]"
    
    except Exception as e:
        return f"[Error extracting text from {file_type}: {str(e)}]"


def extract_quotes_from_text(text: str, client: OpenAI) -> list:
    """Use AI to extract notable quotes from text"""
    if client is None:
        st.error("OpenAI API key is missing. Add it in the sidebar.")
        log_debug("Quote extraction skipped: no API key/client")
        return []
    try:
        log_debug("Extracting quotes from text")
        response = chat_completion(
            client,
            model=get_model_choice(),
            messages=[
                {
                    "role": "system",
                    "content": """You are a senior marketing strategist specializing in B2B SaaS messaging for small business executives.
                    Your mission: Extract the most impactful, quotable statements from Genie AI's Document Compare research that will drive executive-level engagement.

                    CRITICAL IMPACT CRITERIA (rank highest to lowest):
                    1. Quantifiable pain points (e.g., "spend 20 hours", "cost $500/month")
                    2. Emotional frustration statements (e.g., "driving me crazy", "nightmare")
                    3. Executive time/value concerns (e.g., "should be focusing on business")
                    4. Trust/risk statements (e.g., "terrified of missing something")
                    5. Process inefficiency complaints (e.g., "back and forth", "endless revisions")

                    EXTRACT ONLY THE MOST IMPACTFUL statements that would make executives:
                    - Stop scrolling on LinkedIn
                    - Forward to their team
                    - Schedule a demo call
                    - Share with their board

                    Return as JSON with a 'quotes' array of objects, each containing:
                    - 'quote': The punchiest, most memorable version (can be direct or enhanced for impact)
                    - 'source_context': Brief context about the pain point origin
                    - 'impact_score': 1-10 rating of marketing impact potential
                    - 'executive_resonance': Why this resonates with small business leaders
                    Sort quotes highest to lowest by impact_score.

                    Limit to 6-8 most powerful quotes that would dominate marketing content."""
                },
                {
                    "role": "user",
                    "content": f"Extract the most impactful, executive-level quotes from this Document Compare research. Focus on statements that reveal genuine pain points small business leaders face with contract review, document comparison, and legal processes:\n\n{text[:8000]}"
                }
            ],
            response_format={"type": "json_object"}
        )
        
        raw = response.choices[0].message.content
        try:
            result = json.loads(raw)
            quotes = result.get('quotes', [])
            # Sort by impact_score descending if present
            try:
                quotes = sorted(quotes, key=lambda q: q.get('impact_score', 0), reverse=True)
            except Exception:
                pass
            return quotes
        except json.JSONDecodeError:
            # Fallback: return simple list of quotes if JSON parsing fails
            lines = [ln.strip("-• ").strip() for ln in raw.splitlines() if ln.strip()]
            return [{"quote": ln, "source_context": "Uploaded/Source", "impact_score": 6, "executive_resonance": "Relevant pain point"} for ln in lines]
    
    except Exception as e:
        st.error(f"Error extracting quotes: {str(e)}")
        log_debug(f"Quote extraction failed: {e}")
        return []


def generate_draft(client: OpenAI, context: dict) -> str:
    """Generate a content draft based on requirements"""
    if client is None:
        return "Error generating draft: missing OpenAI API key."
    
    # Build the prompt
    source_material = "\n\n---\n\n".join([
        f"**{name}:**\n{text[:3000]}" 
        for name, text in context.get('source_texts', {}).items()
    ])
    
    selected_quotes = "\n".join([
        f'- "{q}"' for q in context.get('selected_quotes', [])
    ])

    additional_context = context.get('additional_context', 'None')
    feedback_notes = context.get('feedback_notes')
    if feedback_notes:
        additional_context += f"\n\nFeedback & Instructions:\n{feedback_notes}"
    
    prompt = f"""Create a {context['content_type']} for {context['platform']} about Genie AI's Document Compare feature.

**Length:** {context['content_length']}
**Feature:** Document Compare - Instantly highlights differences between legal documents
**Tone:** {context['tone']}
**Brand Identity:** Genie AI: values: trust, clarity, simplicity

**Target Audience:** Founders and executives in businesses of 4-50 people, US-based, non-lawyers who want to minimize time and money spent on legal activities.

**Brand Guidelines:**
{context.get('brand_guidelines', 'Professional, trustworthy, helpful. Avoid legal jargon. Focus on time savings and peace of mind.')}

**Example Content Style:**
{context.get('example_content', 'Clear, benefit-focused messaging that speaks to busy executives.')}

**Source Material:**
{source_material}

**Key Quotes to Incorporate:**
{selected_quotes}

**Additional Context:**
{additional_context}

**Workflow Requirements & Constraints:**
{LAUNCH_REQUIREMENTS}

**Audience Profile:**
{AUDIENCE_PROFILE}

**Content Goals:**
1. Generate excitement for Document Compare launch
2. Highlight time savings and reduced legal costs
3. Show how it helps non-lawyers review contracts confidently
4. Include clear call-to-action for trying the feature
5. Match the specified tone and platform requirements
6. Focus on practical benefits over technical features

    Create compelling content that resonates with busy founders and executives who want to spend less time on legal reviews."""

    try:
        log_debug(f"Generating draft for {context['content_type']} ({context['platform']}) length={context['content_length']}")
        response = chat_completion(
            client,
            model=get_model_choice(),
            messages=[
                {
                    "role": "system",
                    "content": """You are an expert marketing content writer. Create engaging, 
                    conversion-focused content that resonates with the target audience.
                    Be creative, compelling, and authentic. Avoid generic marketing speak.
                    You must strictly follow any feedback/instructions provided (e.g., remove emojis if requested). Do not include emojis unless explicitly asked."""
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.8
        )
        return response.choices[0].message.content
    
    except Exception as e:
        log_debug(f"Draft generation failed: {e}")
        return f"Error generating draft: {str(e)}"


def generate_final_content(client: OpenAI, approved_draft: str, context: dict) -> str:
    """Polish and finalize the approved draft"""
    if client is None:
        return "Error finalizing content: missing OpenAI API key."
    
    prompt = f"""Please polish and finalize this approved draft content.

**Approved Draft:**
{approved_draft}

**Content Type:** {context['content_type']}
**Platform:** {context['platform']}
**Tone:** {context['tone']}

**Feedback & Instructions (if any):**
{context.get('feedback_notes', 'None')}

Tasks:
1. Fix any grammar or spelling issues
2. Ensure consistent tone throughout
3. Optimize for the target platform
4. Add any final polish while preserving the approved content structure
5. Ensure all quotes are properly attributed if used

Return the final, publication-ready content."""

    try:
        log_debug(f"Finalizing content for {context['content_type']} ({context['platform']})")
        response = chat_completion(
            client,
            model=get_model_choice(),
            messages=[
                {
                    "role": "system",
                    "content": "You are a senior editor finalizing content for publication. Make minimal changes - only polish and refine. If feedback says remove emojis or follow style rules, enforce them strictly."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content
    
    except Exception as e:
        log_debug(f"Finalization failed: {e}")
        return f"Error finalizing content: {str(e)}"


# ============================================================================
# UI COMPONENTS
# ============================================================================

def render_header():
    """Render the main header"""
    st.markdown("""
    <div class="main-header">
        <h1>📄 GENIE AI Content repurposor and generator</h1>
        <p>AI-powered marketing content for Genie AI's Document Compare feature</p>
    </div>
    """, unsafe_allow_html=True)


def render_progress():
    """Render the step progress indicator"""
    steps = [
        ("📋", "Review Sources"),
        ("🎯", "Content Setup"),
        ("🎨", "Tone & Brand"),
        ("💎", "Extract Quotes"),
        ("📝", "Generate Content"),
        ("✅", "Download All")
    ]
    
    current = st.session_state.current_step
    
    cols = st.columns(len(steps))
    for i, (icon, label) in enumerate(steps):
        step_num = i + 1
        with cols[i]:
            if step_num < current:
                st.markdown(f"<div style='text-align:center;'><span style='font-size:1.5rem;'>✅</span><br><small style='color:#10b981;font-weight:500;'>{label}</small></div>", unsafe_allow_html=True)
            elif step_num == current:
                st.markdown(f"<div style='text-align:center;'><span style='font-size:1.5rem;'>{icon}</span><br><small style='color:#e94560;font-weight:600;'>{label}</small></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align:center;'><span style='font-size:1.5rem;opacity:0.4;'>{icon}</span><br><small style='color:#9ca3af;'>{label}</small></div>", unsafe_allow_html=True)
    
    # Progress bar
    progress = (current - 1) / (len(steps) - 1) * 100
    st.markdown(f"""
    <div class="progress-container">
        <div class="progress-bar" style="width: {progress}%;"></div>
    </div>
    """, unsafe_allow_html=True)


def render_sidebar():
    """Render the sidebar with API key and navigation"""
    with st.sidebar:
        st.markdown("### ⚙️ Settings")
        
        # API Key input
        api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            value=st.session_state.get('openai_api_key', ''),
            help="Enter your OpenAI API key to enable AI features"
        )
        if api_key:
            st.session_state.openai_api_key = api_key
            st.session_state.openai_client = None  # Reset to create new client
        
        # Model selector
        st.markdown("**Model**")
        model_choice = st.selectbox(
            "Select OpenAI model",
            MODEL_OPTIONS,
            index=MODEL_OPTIONS.index(st.session_state.model_choice) if st.session_state.get('model_choice') in MODEL_OPTIONS else 0,
            label_visibility="collapsed",
            help="Used for generation, regeneration, editing, and quote extraction."
        )
        st.session_state.model_choice = model_choice
        
        st.markdown("---")
        
        # Session info
        st.markdown("### 📊 Session Info")
        st.markdown(f"**Files uploaded:** {len(st.session_state.uploaded_files)}")
        st.markdown(f"**Quotes selected:** {len(st.session_state.selected_quotes)}")
        st.markdown(f"**Drafts created:** {len(st.session_state.drafts)}")
        if st.session_state.get("openai_error"):
            st.warning(st.session_state.openai_error)
        
        # Debug log viewer
        with st.expander("🐞 Debug log", expanded=False):
            logs = list(reversed(st.session_state.get("debug_logs", [])))
            st.text_area("Logs", value="\n".join(logs), height=200, label_visibility="collapsed", disabled=True)
        
        st.markdown("---")
        
        # Reset button
        if st.button("🔄 Start Over", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key != 'openai_api_key':
                    del st.session_state[key]
            init_session_state()
            st.rerun()


# ============================================================================
# STEP RENDERERS
# ============================================================================

def render_step_1():
    """Step 1: Review Source Materials"""
    st.markdown("## 📁 Step 1: Document Compare Source Materials")
    st.markdown("Review the pre-loaded source materials for the Document Compare feature launch.")

    # Upload additional files
    st.markdown("### 📤 Upload additional source files")
    st.caption("Add any DOCX, PDF, or TXT files to include in the content generation workflow.")
    uploads = st.file_uploader(
        "Upload files",
        type=["pdf", "docx", "doc", "txt"],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if uploads:
        new_files_added = False
        new_count = 0
        for uploaded in uploads:
            if uploaded.name in st.session_state.uploaded_files:
                continue

            extracted_text = extract_text_from_file(uploaded)

            # ensure unique keys for files with identical names
            base_key = f"upload::{uploaded.name}"
            key = base_key
            suffix = 1
            while key in st.session_state.extracted_texts:
                suffix += 1
                key = f"{base_key} ({suffix})"

                st.session_state.extracted_texts[key] = extracted_text
                st.session_state.uploaded_files.append(uploaded.name)
                new_files_added = True
                new_count += 1

            if new_files_added:
                log_debug(f"Uploaded {new_count} file(s)")
                st.success(f"Added {new_count} uploaded file(s).")
                st.rerun()
    
    # Display pre-loaded source materials
    st.markdown("### 📋 Pre-loaded Source Materials")
    st.info("✅ All Document Compare source materials have been automatically loaded and processed.")
    log_debug("Step 1 rendered; source materials ready.")

    source_files = {
        'product_roadmap': '📊 Product Roadmap Summary',
        'customer_feedback': '💬 Customer Feedback Snippets',
        'engineering_ticket': '🔧 Engineering Ticket (Linear)',
        'meeting_transcript': '🎯 Meeting Transcript',
        'meeting_notes': '📝 Marketing & Product Meeting Notes'
    }

    for key, display_name in source_files.items():
        if key in st.session_state.source_materials:
            text = st.session_state.source_materials[key]
            with st.expander(f"{display_name}", expanded=False):
                if text.startswith("Error"):
                    st.error(text)
                else:
                    st.success(f"✅ Loaded {len(text)} characters")
                    st.text_area(
                        "Content preview",
                        value=text[:1500] + ("..." if len(text) > 1500 else ""),
                        height=150,
                        disabled=True,
                        key=f"source_preview_{key}"
                    )

    # Display uploaded files
    if st.session_state.uploaded_files:
        st.markdown("### 📥 Uploaded Files")
        for key, text in st.session_state.extracted_texts.items():
            if not key.startswith("upload::"):
                continue
            display_name = key.replace("upload::", "")
            with st.expander(f"📄 {display_name}", expanded=False):
                if text.startswith("[Error"):
                    st.error(text)
                else:
                    st.success(f"✅ Loaded {len(text)} characters")
                    st.text_area(
                        "Content preview",
                        value=text[:1500] + ("..." if len(text) > 1500 else ""),
                        height=150,
                        disabled=True,
                        key=f"uploaded_preview_{key}"
                    )
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    
    has_content = len(st.session_state.source_materials) > 0
    
    with col3:
        if st.button("Next: Content Requirements →", type="primary", disabled=not has_content, use_container_width=True):
            log_debug("Navigating to Step 2")
            st.session_state.current_step = 2
            st.rerun()
    
    if not has_content:
        st.warning("⚠️ Source materials not loaded properly. Please check the task-feature-information folder.")
    else:
        st.success("🎯 Ready to proceed with content generation!")


def render_step_2():
    """Step 2: Content requirements for Document Compare launch"""
    st.markdown("## 🎯 Step 2: Launch Content Requirements")
    st.markdown("Plan the launch pieces for Document Compare and set defaults for generation.")
    
    st.markdown("### Launch requirements & audience guardrails")
    st.info(LAUNCH_REQUIREMENTS)
    st.success(AUDIENCE_PROFILE)

    content_types = ["📝 Blog Post", "📱 LinkedIn Post", "📧 Newsletter Email"]
    type_keys = {"📝 Blog Post": "blog", "📱 LinkedIn Post": "linkedin", "📧 Newsletter Email": "newsletter"}

    # Content scope section
    st.markdown("### Content scope (Generate All)")
    st.session_state.content_bundle = st.multiselect(
        "Pieces to auto-generate",
        content_types,
        default=st.session_state.content_bundle or content_types,
        label_visibility="collapsed"
    )
    if not st.session_state.content_bundle:
        st.warning("Select at least one output to generate.")
        st.session_state.content_length = None
    else:
        # set a sensible default content_length based on first selection
        first_key = type_keys[st.session_state.content_bundle[0]]
        st.session_state.content_length = st.session_state.length_preferences[first_key]

    # Content Length preferences per channel
    st.markdown("### Word count preferences (per selected channel)")
    length_options = {
        "blog": [
            "700-900 words (quick skim)",
            "900-1200 words (standard launch story)",
            "1200-1500 words (pillar/SEO depth)"
        ],
        "linkedin": [
            "80-120 words (tight hook + CTA)",
            "120-180 words (standard post)",
            "180-260 words (narrative/story post)"
        ],
        "newsletter": [
            "180-280 words (announcement)",
            "280-400 words (feature highlight)",
            "400-600 words (deep dive update)"
        ]
    }

    selected_keys = [type_keys[t] for t in st.session_state.content_bundle]
    if selected_keys:
        cols = st.columns(len(selected_keys))
        for idx, key in enumerate(selected_keys):
            label = "Blog post" if key == "blog" else "LinkedIn post" if key == "linkedin" else "Newsletter"
            with cols[idx]:
                selection = st.selectbox(
                    label,
                    length_options[key],
                    index=length_options[key].index(st.session_state.length_preferences.get(key, length_options[key][1])) if st.session_state.length_preferences.get(key, length_options[key][1]) in length_options[key] else 1,
                    key=f"len_pref_{key}"
                )
                st.session_state.length_preferences[key] = selection
        # keep a representative content_length for downstream usage
        st.session_state.content_length = st.session_state.length_preferences[selected_keys[0]]
    
    # Feature Focus
    st.markdown("### Feature Focus")
    st.info("📄 **Document Compare** - Instantly highlights differences between legal documents")
    st.session_state.feature_focus = "Document Compare"
    
    # Additional Context
    st.markdown("### Additional Context (Optional)")
    st.markdown("**Target Audience:** Founders and executives in businesses of 4-50 people, US-based, non-lawyers who want to minimize time and money spent on legal activities.")

    additional_context = st.text_area(
        "Additional context",
        value=st.session_state.additional_context,
        placeholder="Any specific messaging, calls-to-action, or additional requirements...",
        height=80,
        label_visibility="collapsed"
    )
    st.session_state.additional_context = additional_context
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 1
            st.rerun()
    with col3:
        if st.button("Next: Tone & Brand →", type="primary", use_container_width=True):
            st.session_state.current_step = 3
            st.rerun()


def render_step_3():
    """Step 3: Tone and brand guidelines"""
    st.markdown("## 🎨 Step 3: Tone & Brand Guidelines")
    
    # Tone Selection
    st.markdown("### Select the tone for your content")
    
    tones = {
        "Professional": "Trustworthy, authoritative, business-focused",
        "Conversational": "Friendly, approachable, executive-to-executive",
        "Practical": "Solution-oriented, results-focused, no-nonsense",
        "Inspiring": "Motivational, empowering, forward-looking",
        "Expert": "Knowledgeable, insightful, industry-aware"
    }
    
    cols = st.columns(3)
    for i, (tone, description) in enumerate(tones.items()):
        with cols[i % 3]:
            selected = st.session_state.tone == tone
            if st.button(
                f"{'✓ ' if selected else ''}{tone}\n\n_{description}_",
                key=f"tone_{tone}",
                use_container_width=True,
                type="primary" if selected else "secondary"
            ):
                st.session_state.tone = tone
                st.rerun()
    
    if st.session_state.tone:
        st.success(f"Selected tone: **{st.session_state.tone}**")
        log_debug(f"Tone selected: {st.session_state.tone}")
    
    st.markdown("---")
    
    # Brand Guidelines
    st.markdown("### Brand Guidelines")
    st.markdown("Add any brand voice guidelines, style rules, or dos/don'ts. Preloaded with a suggested set you can edit.")
    
    brand_guidelines = st.text_area(
        "Brand guidelines",
        value=st.session_state.brand_guidelines,
        placeholder="Edit the suggested guidelines or paste your own.",
        height=200,
        label_visibility="collapsed"
    )
    st.session_state.brand_guidelines = brand_guidelines
    
    st.markdown("---")
    
    # Example Content
    st.markdown("### Example Content (Optional)")
    st.markdown("Paste an example of content you like - we'll match the style.")
    
    example_content = st.text_area(
        "Example content",
        value=st.session_state.example_content,
        placeholder="Paste an example blog post, social post, or any content whose style you'd like to emulate...",
        height=150,
        label_visibility="collapsed"
    )
    st.session_state.example_content = example_content
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 2
            st.rerun()
    with col3:
        if st.button("Next: Select Quotes →", type="primary", disabled=not st.session_state.tone, use_container_width=True):
            st.session_state.current_step = 4
            st.rerun()


def render_step_4():
    """Step 4: Quote extraction and selection"""
    st.markdown("## 💎 Step 4: Select Quotes")
    st.markdown("Choose quotes from your source materials to include in the content.")
    
    client = get_openai_client()
    
    if not client:
        st.warning("⚠️ Please enter your OpenAI API key in the sidebar to extract quotes.")
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("← Back", use_container_width=True):
                st.session_state.current_step = 3
                st.rerun()
        return
    
    # Extract quotes if not already done
    if not st.session_state.extracted_quotes and st.session_state.extracted_texts:
        if st.button("🔍 Extract Quotes from Files", type="primary"):
            with st.spinner("Analyzing files for notable quotes..."):
                all_quotes = []
                for filename, text in st.session_state.extracted_texts.items():
                    quotes = extract_quotes_from_text(text, client)
                    for q in quotes:
                        q['source_file'] = filename
                    all_quotes.extend(quotes)
                all_quotes = sorted(all_quotes, key=lambda q: q.get('impact_score', 0), reverse=True)
                st.session_state.extracted_quotes = all_quotes
                st.rerun()
    
    # Display extracted quotes
    if st.session_state.extracted_quotes:
        st.markdown(f"### Found {len(st.session_state.extracted_quotes)} Quotes")
        
        for i, quote_data in enumerate(st.session_state.extracted_quotes):
            quote_text = quote_data.get('quote', '')
            source = quote_data.get('source_file', 'Unknown')
            context = quote_data.get('source_context', '')
            score = quote_data.get('impact_score', 0)
            
            is_selected = quote_text in st.session_state.selected_quotes
            
            col1, col2 = st.columns([0.1, 0.9])
            with col1:
                if st.checkbox(
                    label=f"Select quote {i+1}",
                    value=is_selected,
                    key=f"quote_{i}",
                    label_visibility="collapsed"
                ):
                    if quote_text not in st.session_state.selected_quotes:
                        st.session_state.selected_quotes.append(quote_text)
                else:
                    if quote_text in st.session_state.selected_quotes:
                        st.session_state.selected_quotes.remove(quote_text)
            
            with col2:
                st.markdown(f"""
                <div class="quote-card">
                    "{quote_text}"
                    <br><small style="color:#92400e;">Source: {source} | Impact: {score}/10</small>
                </div>
                """, unsafe_allow_html=True)
    
    # Manual quote input
    st.markdown("---")
    st.markdown("### Add Custom Quote")
    custom_quote = st.text_input(
        "Enter a custom quote",
        placeholder="Add your own quote to include...",
        label_visibility="collapsed"
    )
    if st.button("Add Quote") and custom_quote:
        st.session_state.selected_quotes.append(custom_quote)
        st.rerun()
    
    # Selected quotes summary
    if st.session_state.selected_quotes:
        st.markdown("---")
        st.markdown(f"### ✅ Selected Quotes ({len(st.session_state.selected_quotes)})")
        for quote in st.session_state.selected_quotes:
            st.markdown(f'- "{quote}"')
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 3
            st.rerun()
    with col3:
        if st.button("Next: Generate Drafts →", type="primary", use_container_width=True):
            st.session_state.current_step = 5
            st.rerun()


def render_step_5():
    """Step 5: Draft generation and review with version history"""
    st.markdown("## 📝 Step 5: Review & Approve Drafts")
    
    client = get_openai_client()
    
    if not client:
        st.warning("⚠️ Please enter your OpenAI API key in the sidebar.")
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("← Back", use_container_width=True):
                st.session_state.current_step = 4
                st.rerun()
        return
    
    # Generate content buttons
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🎯 Generate All Launch Content", type="primary", use_container_width=True):
            with st.status("Generating all launch content pieces...", expanded=True) as status:
                # Generate drafts for the selected formats
                bundle_map = {
                    "📝 Blog Post": {
                        'content_type': '📝 Blog Post',
                        'platform': 'Genie AI Blog',
                        'pref_key': 'blog'
                    },
                    "📱 LinkedIn Post": {
                        'content_type': '📱 LinkedIn Post',
                        'platform': 'Genie AI LinkedIn Page',
                        'pref_key': 'linkedin'
                    },
                    "📧 Newsletter Email": {
                        'content_type': '📧 Newsletter Email',
                        'platform': 'Genie AI Newsletter',
                        'pref_key': 'newsletter'
                    }
                }

                selected_bundle = st.session_state.content_bundle or list(bundle_map.keys())
                content_configs = []
                for label in selected_bundle:
                    if label not in bundle_map:
                        continue
                    entry = bundle_map[label]
                    content_configs.append({
                        'content_type': entry['content_type'],
                        'platform': entry['platform'],
                        'content_length': st.session_state.length_preferences.get(entry['pref_key'], st.session_state.content_length),
                        'tone': st.session_state.tone or 'Professional',
                        'feedback_notes': get_feedback_notes(entry['content_type'])
                    })

                for config in content_configs:
                    status.write(f"Creating {config['content_type']} ({config['content_length']})...")
                    context = {
                        'content_type': config['content_type'],
                        'platform': config['platform'],
                        'content_length': config['content_length'],
                        'feature_focus': 'Document Compare',
                        'tone': config['tone'],
                        'brand_guidelines': st.session_state.brand_guidelines,
                        'example_content': st.session_state.example_content,
                        'source_texts': st.session_state.extracted_texts,
                        'selected_quotes': st.session_state.selected_quotes,
                        'additional_context': st.session_state.additional_context,
                        'feedback_notes': config.get('feedback_notes', '')
                    }

                    draft_content = generate_draft(client, context)

                    new_draft = {
                        'version': len(st.session_state.drafts) + 1,
                        'content': draft_content,
                        'content_type': config['content_type'],
                        'platform': config['platform'],
                        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        'approved': False,
                        'edits': []
                    }

                    st.session_state.drafts.append(new_draft)

                st.session_state.current_draft_index = len(st.session_state.drafts) - 1
                status.update(label="Generation complete", state="complete", expanded=False)
                st.rerun()

    with col2:
        st.markdown("&nbsp;")
    
    with col2:
        if st.session_state.drafts:
            pass
    
    # Version history and per-type tabs
    if st.session_state.drafts:
        st.markdown("---")
        st.markdown("### 📚 Version History")

        tabs = st.tabs(["📱 LinkedIn", "📧 Newsletter", "📝 Blog"])
        tab_map = {
            0: "📱 LinkedIn Post",
            1: "📧 Newsletter Email",
            2: "📝 Blog Post"
        }

        for tab_index, tab in enumerate(tabs):
            ctype = tab_map[tab_index]
            with tab:
                drafts_of_type = [d for d in st.session_state.drafts if d.get('content_type') == ctype]

                if not drafts_of_type:
                    st.info(f"No drafts yet for {ctype}. Generate above.")
                    continue

                # Version pills within tab
                cols = st.columns(min(len(drafts_of_type), 5))
                for i, draft in enumerate(drafts_of_type):
                    with cols[i % 5]:
                        key_idx = st.session_state.current_draft_index_by_type.get(ctype, 0)
                        is_current = i == key_idx
                        label = f"v{draft['version']}" + (" ✓" if draft['approved'] else "")
                        if st.button(label, key=f"{ctype}_v{i}", type="primary" if is_current else "secondary", use_container_width=True):
                            st.session_state.current_draft_index_by_type[ctype] = i
                            st.rerun()

                current_idx = st.session_state.current_draft_index_by_type.get(ctype, 0)
                current_idx = min(current_idx, len(drafts_of_type) - 1)
                st.session_state.current_draft_index_by_type[ctype] = current_idx
                current_draft = drafts_of_type[current_idx]

                st.markdown(f"""
                <div class="draft-header">
                    <span class="version-badge">{ctype}</span>
                    <span class="timestamp">Created: {current_draft['timestamp']}</span>
                </div>
                """, unsafe_allow_html=True)

                length_pref_map = {
                    "📝 Blog Post": st.session_state.length_preferences.get("blog"),
                    "📱 LinkedIn Post": st.session_state.length_preferences.get("linkedin"),
                    "📧 Newsletter Email": st.session_state.length_preferences.get("newsletter"),
                }
                platform_map = {
                    "📝 Blog Post": "Genie AI Blog",
                    "📱 LinkedIn Post": "Genie AI LinkedIn Page",
                    "📧 Newsletter Email": "Genie AI Newsletter",
                }

                edited_content = st.text_area(
                    "Draft content",
                    value=current_draft['content'],
                    height=350,
                    key=f"draft_content_{ctype}_{current_idx}",
                    label_visibility="collapsed"
                )

                if edited_content != current_draft['content']:
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if st.button("💾 Save Changes", key=f"save_{ctype}_{current_idx}"):
                            current_draft['content'] = edited_content
                            current_draft['edits'].append({
                                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                'type': 'manual_edit'
                            })
                            st.success("Changes saved!")
                            st.rerun()

                # Feedback chat for this content type
                st.markdown("#### 💬 Feedback for this piece")
                for msg in st.session_state.feedback_chat.get(ctype, []):
                    with st.chat_message(msg['role']):
                        st.markdown(msg['content'])

                def regenerate_with_feedback():
                    with st.status("Regenerating with feedback...", expanded=True) as status:
                        regen_context = {
                            'content_type': ctype,
                            'platform': platform_map.get(ctype, current_draft.get('platform', 'General')),
                            'content_length': length_pref_map.get(ctype, st.session_state.content_length),
                            'feature_focus': st.session_state.feature_focus,
                            'tone': st.session_state.tone,
                            'brand_guidelines': st.session_state.brand_guidelines,
                            'example_content': st.session_state.example_content,
                            'source_texts': st.session_state.extracted_texts,
                            'selected_quotes': st.session_state.selected_quotes,
                            'additional_context': st.session_state.additional_context,
                            'feedback_notes': get_feedback_notes(ctype)
                        }
                        status.write("Calling OpenAI...")
                        new_content = generate_draft(client, regen_context)
                        status.write("Draft received.")
                        # Overwrite the currently selected draft (respect the version user is viewing)
                        for idx, d in enumerate(st.session_state.drafts):
                            if d is current_draft:
                                st.session_state.drafts[idx]['content'] = new_content
                                st.session_state.drafts[idx]['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                st.session_state.drafts[idx]['approved'] = False
                                st.session_state.drafts[idx].setdefault('edits', []).append({
                                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    'type': 'regen_with_feedback'
                                })
                                break
                        status.update(label="Regeneration complete", state="complete", expanded=False)
                    st.rerun()

                feedback_input = st.chat_input(f"Add feedback for {ctype}", key=f"chat_input_{ctype}")
                if feedback_input:
                    st.session_state.feedback_chat.setdefault(ctype, []).append({"role": "user", "content": feedback_input})
                    st.session_state.feedback_chat[ctype].append({"role": "assistant", "content": "Understood—improving now. Regenerating with your notes."})
                    regenerate_with_feedback()

        # Approval
        st.markdown("---")
        approval = st.checkbox(
            "I approve this draft for final generation",
            value=current_draft['approved'],
            key=f"approval_{ctype}_{current_idx}"
        )
        if approval != current_draft['approved']:
            current_draft['approved'] = approval
            st.session_state.draft_approved = approval
            st.rerun()
        if current_draft['approved']:
            st.success("✅ This draft is approved and ready for final generation!")

        # Regenerate with feedback
        length_pref_map = {
            "📝 Blog Post": st.session_state.length_preferences.get("blog"),
            "📱 LinkedIn Post": st.session_state.length_preferences.get("linkedin"),
            "📧 Newsletter Email": st.session_state.length_preferences.get("newsletter"),
        }
        platform_map = {
            "📝 Blog Post": "Genie AI Blog",
            "📱 LinkedIn Post": "Genie AI LinkedIn Page",
            "📧 Newsletter Email": "Genie AI Newsletter",
        }
        if st.button("♻️ Regenerate with feedback", key=f"regen_{ctype}"):
            with st.status("Regenerating with feedback...", expanded=True) as status:
                regen_context = {
                    'content_type': ctype,
                    'platform': platform_map.get(ctype, current_draft.get('platform', 'General')),
                    'content_length': length_pref_map.get(ctype, st.session_state.content_length),
                    'feature_focus': st.session_state.feature_focus,
                    'tone': st.session_state.tone,
                    'brand_guidelines': st.session_state.brand_guidelines,
                    'example_content': st.session_state.example_content,
                    'source_texts': st.session_state.extracted_texts,
                    'selected_quotes': st.session_state.selected_quotes,
                    'additional_context': st.session_state.additional_context,
                    'feedback_notes': get_feedback_notes(ctype)
                }
                status.write("Calling OpenAI...")
                new_content = generate_draft(client, regen_context)
                status.write("Draft received.")
                new_draft = {
                    'version': len(st.session_state.drafts) + 1,
                    'content': new_content,
                    'content_type': ctype,
                    'platform': regen_context['platform'],
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'approved': False,
                    'edits': [{'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 'type': 'regen_with_feedback'}]
                }
                st.session_state.drafts.append(new_draft)
                st.session_state.current_draft_index_by_type[ctype] = len([d for d in st.session_state.drafts if d.get('content_type') == ctype]) - 1
                status.update(label="Regeneration complete", state="complete", expanded=False)
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 4
            st.rerun()
    with col3:
        has_approved = any(d['approved'] for d in st.session_state.drafts)
        if st.button(
            "Next: Finalize →",
            type="primary",
            disabled=not has_approved,
            use_container_width=True
        ):
            st.session_state.current_step = 6
            st.rerun()
        
        if not has_approved and st.session_state.drafts:
            st.caption("⚠️ Approve a draft to continue")


def render_step_6():
    """Step 6: Final generation and download"""
    st.markdown("## ✅ Step 6: Final Content")
    
    client = get_openai_client()
    if not client:
        st.warning("⚠️ Please enter your OpenAI API key in the sidebar to finalize and download content.")
        if st.button("← Back to Drafts"):
            st.session_state.current_step = 5
            st.rerun()
        return
    
    # Get all approved drafts; if a type has none approved, take the latest draft so everything is pulled
    approved_drafts = [d for d in st.session_state.drafts if d['approved']]
    drafts_by_type = {}
    for d in st.session_state.drafts:
        ctype = d.get('content_type')
        if ctype not in drafts_by_type or d['version'] > drafts_by_type[ctype]['version']:
            drafts_by_type[ctype] = d
    for ctype, latest in drafts_by_type.items():
        if not any(d.get('content_type') == ctype for d in approved_drafts):
            approved_drafts.append(latest)
    
    if not approved_drafts:
        st.warning("No drafts available. Please generate content first.")
        if st.button("← Back to Drafts"):
            st.session_state.current_step = 5
            st.rerun()
        return
    
    # Summary of approved content
    st.markdown("### 📋 Launch Content Summary")

    st.markdown(f"**Feature:** {st.session_state.feature_focus}")
    st.markdown(f"**Tone:** {st.session_state.tone}")
    st.markdown(f"**Approved Content Pieces:** {len(approved_drafts)}")

    for draft in approved_drafts:
        st.markdown(f"- ✅ {draft.get('content_type', 'Content')} (v{draft['version']})")
    
    st.markdown("---")
    
    # Generate final versions for all approved drafts
    if not st.session_state.get('final_contents'):
        st.session_state.final_contents = {}

    # Auto-finalize all approved drafts on entry
    for draft in approved_drafts:
        draft_id = f"{draft['content_type']}_v{draft['version']}"
        if draft_id not in st.session_state.final_contents:
            with st.spinner(f"Polishing {draft['content_type']} (v{draft['version']})..."):
                context = {
                    'content_type': draft['content_type'],
                    'platform': draft.get('platform', 'General'),
                    'tone': st.session_state.tone,
                    'feedback_notes': get_feedback_notes(draft.get('content_type', ''))
                }
                final = generate_final_content(client, draft['content'], context)
                st.session_state.final_contents[draft_id] = final
    
    # Display all final content pieces
    if st.session_state.final_contents:
        st.markdown("### 🎉 Final Launch Content")
        
        st.markdown(f"""
        <div class="approval-section">
            <p style="margin:0;color:#065f46;font-weight:500;">✅ All content pieces are ready for download!</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Display each content piece
        for draft in approved_drafts:
            draft_id = f"{draft['content_type']}_v{draft['version']}"
            if draft_id in st.session_state.final_contents:
                final_content = st.session_state.final_contents[draft_id]

                with st.expander(f"📄 {draft['content_type']} - Final Version", expanded=True):
                    st.text_area(
                        f"Final {draft['content_type']} content",
                        value=final_content,
                        height=300,
                        label_visibility="collapsed",
                        key=f"final_display_{draft_id}"
                    )

                    # Download buttons for this piece
                    col1, col2, col3 = st.columns(3)

                    content_filename = draft['content_type'].replace(' ', '_').replace('📝', 'Blog').replace('📱', 'LinkedIn').replace('📧', 'Newsletter')
        
                    with col1:
                        st.download_button(
                            label="📄 TXT",
                            data=final_content,
                            file_name=f"{content_filename}_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            use_container_width=True,
                            key=f"download_txt_{draft_id}"
                        )
        
                    with col2:
                        st.download_button(
                            label="📝 Markdown",
                            data=final_content,
                            file_name=f"{content_filename}_{datetime.now().strftime('%Y%m%d')}.md",
                            mime="text/markdown",
                            use_container_width=True,
                            key=f"download_md_{draft_id}"
                        )
        
                    with col3:
                        # Create JSON export with metadata
                        export_data = {
                            'content': final_content,
                            'metadata': {
                                'content_type': draft['content_type'],
                                'platform': draft.get('platform', 'General'),
                                'tone': st.session_state.tone,
                                'feature_focus': st.session_state.feature_focus,
                                'generated_at': datetime.now().isoformat(),
                                'quotes_used': st.session_state.selected_quotes,
                                'source_files': list(st.session_state.extracted_texts.keys())
                            }
                        }
                        
                        st.download_button(
                            label="📦 JSON",
                            data=json.dumps(export_data, indent=2),
                            file_name=f"{content_filename}_{datetime.now().strftime('%Y%m%d')}.json",
                            mime="application/json",
                            use_container_width=True,
                            key=f"download_json_{draft_id}"
                        )
        
        # Removed bulk/launch pack downloads to keep UI simpler
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("← Back to Drafts", use_container_width=True):
            st.session_state.current_step = 5
            st.rerun()
    with col3:
        if st.button("🔄 Create New Content", use_container_width=True):
            # Keep API key, reset everything else
            api_key = st.session_state.get('openai_api_key')
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            init_session_state()
            if api_key:
                st.session_state.openai_api_key = api_key
            st.rerun()


# ============================================================================
# MAIN APP
# ============================================================================

def main():
    render_header()
    render_sidebar()
    render_progress()
    
    st.markdown("---")
    
    # Route to current step
    step_renderers = {
        1: render_step_1,
        2: render_step_2,
        3: render_step_3,
        4: render_step_4,
        5: render_step_5,
        6: render_step_6
    }
    
    renderer = step_renderers.get(st.session_state.current_step, render_step_1)
    renderer()


if __name__ == "__main__":
    main()
